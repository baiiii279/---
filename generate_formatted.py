#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
根据 temp_template.docx 的格式，
将 PaperCraft_第1-2章.docx 的内容重新排版输出。
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy
import re


def set_run_font(run, font_cn='宋体', font_en='Times New Roman', size=Pt(12), bold=None):
    """设置 run 的中英文字体"""
    run.font.size = size
    run.font.bold = bold
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn('w:eastAsia'), font_cn)


def add_paragraph_with_style(doc, text, style_name, font_cn='宋体', font_en='Times New Roman',
                              size=None, bold=None, alignment=None, first_indent=None,
                              line_spacing=None):
    """添加段落到文档"""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)

    p.alignment = alignment

    if first_indent is not None:
        p.paragraph_format.first_line_indent = first_indent

    set_run_font(run, font_cn, font_en, size, bold)
    return p, run


def add_title_heading(doc, text, level=1):
    """添加标题段落"""
    if level == 1:
        style = '1标题'
        p, _ = add_paragraph_with_style(doc, text, style,
                                         font_cn='黑体', size=Pt(15),
                                         alignment=WD_ALIGN_PARAGRAPH.CENTER)
    elif level == 2:
        style = '2标题'
        p, _ = add_paragraph_with_style(doc, text, style,
                                         font_cn='黑体', size=Pt(14))
    else:
        style = '3标题'
        p, _ = add_paragraph_with_style(doc, text, style,
                                         font_cn='黑体', size=Pt(12))
    return p


def add_body_text(doc, text, bold_first=False):
    """添加正文段落"""
    style = '正文1'
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(12), bold=None)
    return p


def add_body_text_with_bold(doc, parts):
    """添加正文段落，支持部分文字加粗
    parts: list of (text, is_bold)
    """
    style = '正文1'
    p = doc.add_paragraph(style=style)
    for text, is_bold in parts:
        run = p.add_run(text)
        set_run_font(run, '宋体', 'Times New Roman', Pt(12), is_bold)
    return p


def add_caption(doc, text):
    """添加图表标题"""
    style = 'Caption'
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', Pt(10.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_empty_para(doc, style='Normal'):
    """添加空段落"""
    p = doc.add_paragraph(style=style)
    return p


def add_centered_text(doc, text, size=Pt(15), bold=True, font_cn='宋体'):
    """添加居中文本（封面用）"""
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, font_cn, 'Times New Roman', size, bold)
    return p


def build_cover(doc):
    """构建封面"""
    # 空行 - 用大号空白占位
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('')
    run.font.size = Pt(26)

    # 主标题
    add_centered_text(doc, '教学实践Ⅱ：软件项目开发课程设计', Pt(15), True)

    # 课程教师
    add_centered_text(doc, '（课程教师：黄炜钦）', Pt(15), True)

    # 空行
    add_empty_para(doc)

    # 项目标题
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(2.2)
    run = p.add_run('题   目：PaperCraft — 多智能体论文写作助手')
    set_run_font(run, '黑体', 'Times New Roman', Pt(18), True)

    # 空行
    add_empty_para(doc)

    # 成员表
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(2.2)
    run = p.add_run('成员表')
    set_run_font(run, '宋体', 'Times New Roman', Pt(18), True)

    # 注
    p = doc.add_paragraph(style='Normal')
    run = p.add_run('注：该表用于衡量个人工作量，请务必认真填写。')
    set_run_font(run, '宋体', 'Times New Roman', Pt(12), True)

    # 日期
    add_centered_text(doc, '2026年6月29日', Pt(15), False, '宋体')

    # 空行
    add_empty_para(doc)
    add_empty_para(doc)

    # 目录
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目    录')
    set_run_font(run, '黑体', 'Times New Roman', Pt(14))

    # 目录自动生成提示
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('（完成全部内容后，在Word中使用"引用→目录"自动生成）')
    set_run_font(run, '宋体', 'Times New Roman', Pt(12))

    # 分页 - 目录后分页
    doc.add_page_break()


def build_chapter1(doc):
    """构建第1章"""
    add_title_heading(doc, '第1章  开发环境与技术', 1)

    add_title_heading(doc, '1.1  开发环境与工具', 2)

    add_body_text(doc, '本系统使用Visual Studio Code作为主要代码编辑器，支持Python、JavaScript、TypeScript等多种语言开发。后端使用Python 3.10开发，前端使用Node.js 18作为运行环境，数据库采用MySQL 8.0进行数据存储。版本控制使用Git，API接口测试使用Postman。')

    add_title_heading(doc, '1.2  相关技术', 2)

    add_body_text(doc, '本系统采用前后端分离架构，主要技术包括：')

    # 技术列表 - 每项用加粗的技术名+正文描述
    techs = [
        ('FastAPI', '（0.115+）：基于Python的现代Web框架，用于构建后端RESTful API接口，并利用其异步特性实现SSE实时推送Agent执行状态。'),
        ('React', ' 18：由Meta开发的JavaScript前端框架，采用组件化开发模式构建单页应用，配合React Router实现页面路由管理，使用TypeScript进行类型检查。'),
        ('Vite', ' 5：现代化前端构建工具，利用浏览器原生ES模块导入实现毫秒级热更新开发体验。'),
        ('CrewAI', '（0.76+）：多智能体编排框架，本系统使用其编排5个Agent（文献解析、大纲生成、内容撰写、润色、引用检查）实现端到端的论文写作自动化协作。'),
        ('DeepSeek API', '：国内大语言模型API，为每个Agent提供自然语言理解和生成能力，性价比高且中文表现优秀。'),
        ('SQLAlchemy', '（2.0+）：Python ORM框架，通过对象关系映射操作MySQL数据库，定义了6张数据表。'),
        ('JWT', '：JSON Web Token，用于前后端分离架构下的用户认证，配合bcrypt算法保障密码安全。'),
    ]

    for name, desc in techs:
        p = doc.add_paragraph(style='正文1')
        run1 = p.add_run(name)
        set_run_font(run1, '宋体', 'Times New Roman', Pt(12), True)
        run2 = p.add_run(desc)
        set_run_font(run2, '宋体', 'Times New Roman', Pt(12))

    # 第1章结束，分页
    doc.add_page_break()


def build_chapter2(doc):
    """构建第2章"""
    add_title_heading(doc, '第2章  系统需求分析', 1)

    # 2.1
    add_title_heading(doc, '2.1  功能需求分析', 2)

    add_body_text(doc, '本系统——PaperCraft多智能体论文写作助手——是一个基于Multi-Agent架构的智能论文写作协作平台，主要面向需要撰写学术论文的高校学生和研究人员。系统采用前后端分离架构，支持用户注册登录、论文管理、Agent辅助写作、文献管理等功能。系统参与者分为普通用户和管理员两种角色，各角色的权限说明如表2-1所示。')

    # 表2-1
    add_caption(doc, '表2-1  系统参与者权限表')
    # 这里预留表格位置

    add_body_text(doc, '为了直观地展示各种角色对应的功能及其之间的联系，采用系统用例图进行描述。系统总体用例图如图2-1所示。')

    # 图2-1
    add_caption(doc, '图2-1  系统总体用例图（此处应插入用例图）')

    add_body_text(doc, '系统总体用例图包含两个角色：用户和管理员。用户的用例包括：注册与登录、创建论文、管理文献库、启动Agent写作流程、审阅Agent产出、查看个人信息、修改密码等。管理员的用例包括：用户管理、论文管理、查看系统日志、查看系统统计等。')

    # 2.1.1
    add_title_heading(doc, '2.1.1  用户功能需求', 3)

    user_items = [
        ('（1）注册与登录', '用户可以通过注册功能创建账号，注册时需要提供用户名和密码，系统对用户名唯一性、密码长度（不少于6位）进行校验。登录时验证用户名和密码的正确性，验证通过后返回JWT令牌用于后续请求的身份认证。登录注册界面需提供明确的错误提示信息，如"用户名已存在""用户名或密码错误"等。同时，前端需做好表单验证，包括输入格式校验、必填项检查等，提升用户体验。'),
        ('（2）个人中心管理', '用户可以在个人中心页面查看和编辑个人信息，包括上传头像、修改邮箱、修改密码等。个人中心还展示用户已创建的论文列表，并提供"我的文献库"（即我的收藏）的快捷入口，方便用户快速查找和管理自己的学术资源。修改密码时需要验证原密码正确性，确保账号安全。'),
        ('（3）论文写作管理', '用户创建论文时需要输入论文主题、选择论文模板（课程论文模板或期刊论文模板），并可从个人文献库中勾选参考文献进行关联。创建后可查看论文列表，列表中展示论文标题、当前状态（如"草稿""文献解析中""大纲生成中""内容撰写中""已完成"等）和创建时间。用户可点击进入论文工作台进行详细的Agent写作操作。'),
        ('（4）Agent辅助写作', 'Agent辅助写作是本系统的核心功能，采用Human-in-the-loop（人在回路）的协作模式。用户在每个Agent执行完成后可以审阅产出内容，选择确认、驳回（附修改意见）或直接编辑修改。具体流程如下：'),
    ]

    for title, desc in user_items:
        p = doc.add_paragraph(style='正文1')
        run1 = p.add_run(title)
        set_run_font(run1, '宋体', 'Times New Roman', Pt(12))
        run2 = p.add_run('\n' + desc)
        set_run_font(run2, '宋体', 'Times New Roman', Pt(12))

    # Agent 分阶段描述
    stages = [
        '文献解析阶段：用户提供参考文献后，文献解析Agent自动提取每篇文献的研究问题、方法和结论，输出结构化文献摘要。用户确认或驳回后进入下一步。',
        '大纲生成阶段：大纲生成Agent根据已确认的文献摘要和论文模板，生成包含章节标题和写作要点的论文大纲。用户可调整章节顺序、增删要点。',
        '内容撰写阶段：内容撰写Agent按照已确认的大纲逐章撰写论文内容，支持逐章确认或全量确认。',
        '润色优化阶段：润色Agent对论文进行语言流畅度提升、重复表达改写、术语统一等优化。',
        '引用检查阶段：引用检查Agent校验论文中所有引用的准确性和格式规范性，输出检查报告。',
    ]

    for stage in stages:
        p = doc.add_paragraph(style='正文1')
        # 使阶段名称加粗
        colon_idx = stage.index('：')
        name = stage[:colon_idx+1]
        desc = stage[colon_idx+1:]
        run1 = p.add_run(name)
        set_run_font(run1, '宋体', 'Times New Roman', Pt(12), True)
        run2 = p.add_run(desc)
        set_run_font(run2, '宋体', 'Times New Roman', Pt(12))

    # (5) 文献库管理
    p = doc.add_paragraph(style='正文1')
    run1 = p.add_run('（5）文献库管理')
    set_run_font(run1, '宋体', 'Times New Roman', Pt(12))
    run2 = p.add_run('\n用户可以在"我的文献库"页面管理个人学术文献，支持新增文献（填写标题、作者、出处、摘要等信息）、查看文献详情、删除文献等操作。文献库中的文献可在创建论文时被勾选引用，实现文献的跨论文复用。这一功能同时满足了课程设计对"个人页面-我的收藏"功能的要求。')
    set_run_font(run2, '宋体', 'Times New Roman', Pt(12))

    # 2.1.2
    add_title_heading(doc, '2.1.2  管理员功能需求', 3)

    admin_items = [
        ('（1）用户管理', '管理员可以查看系统所有用户列表，了解用户注册情况和账号信息，必要时可管理用户状态。'),
        ('（2）论文管理', '管理员可以查看系统所有论文列表，按状态筛选论文，了解系统使用情况，必要时可删除违规或异常的论文。'),
        ('（3）系统日志查看', '管理员可以查看Agent执行的详细日志，包括每个Agent的输入内容、输出内容、执行耗时、成功失败状态等，用于监控系统运行状况和排查问题。日志支持按用户、论文、状态等多维度筛选。'),
        ('（4）系统统计', '管理员可以查看系统基础统计数据，如用户总数、论文总数等，了解系统整体使用情况。'),
    ]

    for title, desc in admin_items:
        p = doc.add_paragraph(style='正文1')
        run1 = p.add_run(title)
        set_run_font(run1, '宋体', 'Times New Roman', Pt(12))
        run2 = p.add_run('\n' + desc)
        set_run_font(run2, '宋体', 'Times New Roman', Pt(12))

    # 2.1.3
    add_title_heading(doc, '2.1.3  Agent协作流程', 3)

    add_body_text(doc, '本系统设计了5个AI Agent协同完成论文写作任务，各Agent的具体职责如表2-2所示。')

    # 表2-2
    add_caption(doc, '表2-2  Agent职责表')

    add_body_text(doc, 'Agent协作流程如图2-2所示。用户输入论文主题并提交参考文献后，Orchestrator（编排管理器）按顺序依次调度各个Agent：文献解析Agent首先执行，输出结构化文献摘要供用户确认；确认通过后，大纲生成Agent根据文献摘要和模板生成论文大纲；用户确认大纲后，内容撰写Agent逐章撰写论文；撰写完成后由润色Agent进行语言优化；最后引用检查Agent校验引用准确性。每个步骤完成后系统都会等待用户反馈，实现Human-in-the-loop协作。')

    # 图2-2
    add_caption(doc, '图2-2  Agent协作流程图（此处应插入流程图）')

    # 2.2
    add_title_heading(doc, '2.2  非功能需求分析', 2)

    # 2.2.1
    add_title_heading(doc, '2.2.1  性能需求', 3)
    add_body_text(doc, '性能需求方面，页面响应时间应控制在2秒以内，确保用户操作的流畅性。单个Agent执行应在30秒内完成（取决于大语言模型API响应速度），若超时则向用户提示异常并提供重试选项。系统应支持至少10个用户同时使用，确保多用户并发场景下的稳定性。Agent执行过程中，通过SSE（Server-Sent Events）技术实时向前端推送状态变更，用户可随时了解Agent执行进度，避免长时间等待的无反馈问题。')

    # 2.2.2
    add_title_heading(doc, '2.2.2  安全需求', 3)
    add_body_text(doc, '安全需求方面，用户密码使用bcrypt算法进行哈希存储，不存储明文密码，防止密码泄露风险。用户认证采用JWT令牌机制，令牌有效期为7天，过期后需重新登录。前后端通信均需携带JWT令牌，后端对每个需要进行身份验证的API请求进行令牌验证和解析。管理员接口需校验用户角色权限，防止普通用户越权访问管理功能。用户注册时需对用户名唯一性进行校验，防止重复注册。')

    # 2.2.3
    add_title_heading(doc, '2.2.3  可用性需求', 3)
    add_body_text(doc, '可用性需求方面，界面风格采用编辑主义×瑞士风格的设计语言，以深蓝黑（#0F172A）为主色、亮蓝（#2563EB）为点缀色，背景使用极浅灰蓝（#F8FAFC），整体风格简洁专业。标题使用思源宋体（Noto Serif SC），正文使用思源黑体（Noto Sans SC），字体层级清晰。页面布局采用非对称两栏设计和慷慨的间距，卡片使用6px小圆角和细边框分割，无阴影。Agent写作工作台采用横向流程管道视图，直观展示5个流程节点，当前活跃节点高亮显示。所有用户操作界面提供清晰的反馈和错误提示，包括表单验证错误、Agent执行异常提示等。')

    # 2.2.4
    add_title_heading(doc, '2.2.4  可维护性需求', 3)
    add_body_text(doc, '可维护性需求方面，后端采用模块化架构，按功能模块划分路由，每个模块职责清晰。Agent层封装统一基类（BaseAgent），新增Agent只需继承基类并实现run方法，具有良好的可扩展性。所有Agent执行日志记录到数据库的agent_logs表中，包括每一步的输入输出、执行耗时和状态，便于问题追溯和系统调试。代码遵循PEP 8编码规范，关键函数添加类型提示，前后端代码均使用TypeScript或Python类型注解，提高代码可读性和健壮性。')


def main():
    # 打开模板并读取边距设置
    template_doc = Document('temp_template.docx')
    template_section = template_doc.sections[0]

    # 记录模板的边距和页面设置
    page_width = template_section.page_width
    page_height = template_section.page_height
    top_margin = template_section.top_margin
    bottom_margin = template_section.bottom_margin
    left_margin = template_section.left_margin
    right_margin = template_section.right_margin

    doc = template_doc

    # 获取 body 元素
    body = doc.element.body

    # 只移除段落元素，保留 section properties
    for child in list(body):
        tag = child.tag
        if tag.endswith('}sectPr'):
            continue
        if tag.endswith('}p'):
            body.remove(child)

    # 确保 section 边距与模板一致
    for section in doc.sections:
        section.page_width = page_width
        section.page_height = page_height
        section.top_margin = top_margin
        section.bottom_margin = bottom_margin
        section.left_margin = left_margin
        section.right_margin = right_margin

    # 构建内容
    build_cover(doc)
    build_chapter1(doc)
    build_chapter2(doc)

    # 确保最终边距一致
    for section in doc.sections:
        section.top_margin = top_margin
        section.bottom_margin = bottom_margin
        section.left_margin = left_margin
        section.right_margin = right_margin

    # 存储
    output_path = 'PaperCraft_第1-2章_格式化后.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')


if __name__ == '__main__':
    main()
