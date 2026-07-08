#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成 PaperCraft 课程设计论文 第3章 系统设计 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_run_font(run, font_cn='宋体', font_en='Times New Roman', size=Pt(12), bold=None):
    """设置 run 的中英文字体"""
    run.font.size = size
    if bold is not None:
        run.font.bold = bold
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn('w:eastAsia'), font_cn)


def set_paragraph_spacing(paragraph, line_spacing=20):
    """设置段落行距固定值20磅"""
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(line_spacing * 20))  # 20磅 = 400
    spacing.set(qn('w:lineRule'), 'exact')


def add_body_paragraph(doc, text, first_indent=True, font_size=Pt(12)):
    """添加正文段落（小四号，宋体，首行缩进2字符）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_cn='宋体', font_en='Times New Roman', size=font_size)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(24)  # 约2字符
    set_paragraph_spacing(p)
    return p


def add_heading_para(doc, text, level=1, font_size=Pt(14)):
    """添加标题段落（黑体）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        # 一级标题：四号（14pt），居中
        set_run_font(run, font_cn='黑体', font_en='Times New Roman', size=Pt(14), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        # 二级标题：小四（12pt），左对齐
        set_run_font(run, font_cn='黑体', font_en='Times New Roman', size=Pt(12), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 3:
        # 三级标题：小四（12pt），左对齐
        set_run_font(run, font_cn='黑体', font_en='Times New Roman', size=Pt(12), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(24)
    set_paragraph_spacing(p)
    return p


def add_figure_caption(doc, text):
    """添加图注（五号，居中）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_cn='宋体', font_en='Times New Roman', size=Pt(10.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p)
    return p


def add_table_caption(doc, text):
    """添加表标题（五号，居中，加粗）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_cn='宋体', font_en='Times New Roman', size=Pt(10.5), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p)
    return p


def add_image_placeholder(doc, label):
    """添加图片占位框"""
    p = doc.add_paragraph()
    run = p.add_run(f'[此处插入{label}]')
    set_run_font(run, font_cn='宋体', font_en='Times New Roman', size=Pt(10.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.color.rgb = RGBColor(128, 128, 128)
    set_paragraph_spacing(p)
    return p


def create_chapter3():
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')

    # ========== 第3章 系统设计 ==========
    add_heading_para(doc, '第3章 系统设计', level=1)

    # 3.1
    add_heading_para(doc, '3.1 系统总体设计', level=2)

    add_body_paragraph(doc, '系统总体设计面向全局问题，从宏观层面确定系统的整体架构、模块划分及相互关系，为后续详细设计与编码实现提供蓝图。本系统——PaperCraft多智能体论文写作助手——采用"React前端单页应用 + FastAPI后端服务 + CrewAI多智能体编排层 + MySQL数据库"的分层架构。其中，用户通过浏览器访问React构建的SPA前端界面；前端通过HTTP/HTTPS协议调用后端RESTful API，并通过SSE（Server-Sent Events）接收Agent执行过程中的实时流式输出；后端基于FastAPI框架提供统一的API接口，利用其异步特性实现高效的请求处理和SSE推送；CrewAI作为多智能体编排框架，管理5个Agent的协作流程；所有业务数据持久化存储于MySQL数据库中；AI能力通过DeepSeek API提供，后端通过litellm库统一封装各Agent的语言模型调用。整体架构遵循典型的分层结构：表现层（React SPA）、业务逻辑层（FastAPI API路由）、Agent编排层（CrewAI + 5个Agent）、数据访问层（SQLAlchemy ORM）和数据存储层（MySQL）。')

    add_body_paragraph(doc, '系统各层职责明确：表现层负责用户交互和界面渲染，采用React 18框架配合React Router实现前端路由管理，使用TypeScript进行类型检查，Vite 5作为构建工具提供毫秒级热更新开发体验；业务逻辑层负责处理HTTP请求、参数校验、权限验证和业务数据流转，基于FastAPI（0.115+）实现RESTful API，所有需要身份验证的接口通过JWT令牌进行鉴权；Agent编排层负责管理多个AI Agent的协作执行，基于CrewAI（0.76+）框架调度5个Agent（文献解析Agent、大纲生成Agent、内容撰写Agent、润色Agent、引用检查Agent），每个Agent封装在独立模块中并继承统一基类；数据访问层使用SQLAlchemy 2.0+ ORM框架，通过对象关系映射操作MySQL数据库，定义了6张ORM模型表。系统架构图如图3-1所示。')

    add_image_placeholder(doc, '图3-1 系统架构图')

    add_body_paragraph(doc, '在前端层面，系统被划分为8个主要页面，分别是首页、登录页、注册页、论文列表页、论文工作台页、文献库页、个人中心页和管理后台页。首页展示系统介绍和功能入口；登录与注册页负责用户身份认证；论文列表页展示用户创建的所有论文及其状态；论文工作台页为系统的核心页面，集成Agent写作流水线UI和流式输出面板；文献库页提供文献的增删改查功能；个人中心页展示用户信息和密码修改；管理后台页为管理员提供用户管理、论文管理和系统日志查看等功能。')

    add_body_paragraph(doc, '在后端层面，系统按功能划分为认证管理（/api/auth）、用户管理（/api/user）、论文管理（/api/papers）、Agent执行（/api/papers/{id}/agent）、文献管理（/api/user/references）和管理员管理（/api/admin）六大模块。每个模块对应独立的路由文件，实现关注点分离。其中Agent执行模块既支持同步调用也支持异步调用，异步模式下通过BackgroundTasks + SSE机制实现流式输出，前端通过EventSource接收实时事件。')

    add_body_paragraph(doc, '系统功能结构如图3-2所示，从用户维度和管理员维度展示了系统的功能模块划分。')

    add_image_placeholder(doc, '图3-2 系统功能结构图')

    # 3.2
    add_heading_para(doc, '3.2 系统详细设计', level=2)

    add_body_paragraph(doc, '系统详细设计针对每个核心功能模块进行具体的流程设计，明确各模块的执行步骤、数据流转和异常处理逻辑。本系统根据功能需求分析的结果，将详细设计分为注册与登录模块、论文管理模块、Agent辅助写作模块、文献库管理模块、个人中心模块和管理后台模块六个部分进行描述。')

    # 3.2.1
    add_heading_para(doc, '3.2.1 注册与登录模块', level=2)

    add_body_paragraph(doc, '注册与登录功能是用户使用系统的基础功能。用户进入系统后，若未登录会被重定向至登录页面。系统会判断用户是否已经拥有账号。如果用户已有账号，则输入用户名和密码进行登录，系统对登录信息进行校验。当用户输入的信息正确时，系统生成JWT令牌并返回登录成功信息，前端将令牌存入localStorage并跳转至系统首页；当用户输入的账号或密码不正确时，系统会提示"用户名或密码错误"，用户需要重新输入登录信息。')

    add_body_paragraph(doc, '如果用户没有账号，则可以点击"注册"链接进入注册页面。注册时，用户需要填写用户名、邮箱和密码等基本信息。系统会对用户填写的信息进行校验，检查用户名是否已存在、密码长度是否不少于6位、邮箱格式是否正确等。如果用户名已经存在，系统会提示"用户名已存在"；如果注册信息校验通过，则系统使用bcrypt算法对密码进行哈希加密后写入数据库，完成用户注册，并提示"注册成功"。注册成功后，用户可以返回登录页面，使用新注册的账号进行登录。登录与注册程序流程图如图3-3所示。')

    add_image_placeholder(doc, '图3-3 登录与注册程序流程图')

    add_body_paragraph(doc, '登录模块的时序交互过程为：用户在登录页面输入用户名和密码，点击登录按钮后，前端首先进行表单校验（验证用户名和密码是否为空），校验通过后调用后端POST /api/auth/login接口。后端收到请求后查询用户表验证用户名是否存在，使用bcrypt验证密码哈希，验证通过后生成JWT令牌（有效期为7天），返回用户信息和token。前端收到响应后将token存入localStorage，并跳转至系统首页。登录时序图如图3-4所示。')

    add_image_placeholder(doc, '图3-4 登录时序图')

    # 3.2.2
    add_heading_para(doc, '3.2.2 论文管理模块', level=2)

    add_body_paragraph(doc, '论文管理模块是系统的核心业务模块之一，负责论文的创建、查看、编辑和删除等操作。用户进入论文列表页面后，系统会读取当前用户的所有论文，并按照创建时间倒序展示。列表中的每篇论文展示其标题、当前状态（如"草稿""文献解析中""大纲生成中""已完成"等）和创建时间，不同状态使用不同颜色标签进行区分，便于用户快速了解论文进展。')

    add_body_paragraph(doc, '用户点击"新建论文"按钮时，弹出创建论文模态框。用户需要输入论文主题、选择论文模板（课程论文模板或期刊论文模板），并可从个人文献库中勾选参考文献。用户确认创建后，系统调用后端POST /api/papers接口，后端将论文记录插入papers表，同时将选中的文献关联记录插入paper_references表。创建成功后，论文出现在论文列表中，初始状态为"草稿"。')

    add_body_paragraph(doc, '用户点击某篇论文即可进入论文工作台页面，查看该论文的详细信息并进行Agent写作操作。在论文工作台中，页面会展示论文的完整内容（包含已生成的各章节内容），同时提供Agent流水线操作界面。创建论文程序流程图如图3-5所示。')

    add_image_placeholder(doc, '图3-5 创建论文程序流程图')

    # 3.2.3
    add_heading_para(doc, '3.2.3 Agent辅助写作模块', level=2)

    add_body_paragraph(doc, 'Agent辅助写作模块是本系统的核心创新模块，采用Human-in-the-loop（人在回路）的协作模式。系统按照"文献解析（parse）→大纲生成（outline）→内容撰写（write）→润色优化（polish）→引用检查（cite_check）"的顺序依次调度5个Agent，每个Agent执行完成后等待用户审阅和确认，用户可以选择批准（approve）、驳回（reject，附修改意见）或直接编辑（edit）修改产出内容。论文状态严格遵循draft→parsing→outlining→writing→polishing→checking→complete的七阶段流转顺序。')

    add_body_paragraph(doc, 'Agent辅助写作流程如下：用户首先在论文工作台中点击"开始写作"按钮触发全流程。系统首先调度文献解析Agent，该Agent读取用户关联的参考文献，利用大语言模型解析每篇文献的研究问题、方法和结论，输出结构化文献摘要。用户审阅摘要内容后，可选择"确认"进入下一阶段，或"驳回"并附上修改意见让Agent重新生成。')

    add_body_paragraph(doc, '文献解析确认通过后，系统调度大纲生成Agent。该Agent根据已确认的文献摘要和用户选择的论文模板，生成包含章节标题和写作要点的论文大纲（JSON格式）。用户审阅大纲后同样可以确认或驳回。审阅通过后，系统进入内容撰写阶段。')

    add_body_paragraph(doc, '内容撰写Agent按照已确认的大纲逐章撰写论文内容，输出Markdown格式的完整章节。用户可逐章确认或一次性确认全部内容。确认后进入润色阶段，润色Agent对论文进行语言流畅度提升、重复表达改写和术语统一等优化。最后，引用检查Agent校验论文中所有引用的准确性和格式规范性，输出引用检查报告。引用检查通过后，论文状态更新为"已完成"。')

    add_body_paragraph(doc, 'Agent辅助写作流程图如图3-6所示，展示了从启动写作到论文完成的全流程交互过程。')

    add_image_placeholder(doc, '图3-6 Agent辅助写作流程图')

    add_body_paragraph(doc, 'Agent工作流的状态转换关系如图3-7所示。每个Agent执行单元具有明确的五阶段状态：pending（待执行）→ running（执行中）→ 执行完成后进入等待用户反馈状态，用户可选择approve（批准）、reject（驳回）或edit（编辑）。批准后进入下一个Agent流程，驳回后Agent根据反馈重新生成，编辑则允许用户直接修改内容。任一Agent执行失败（error状态）时，系统提示用户重试或终止流程。')

    add_image_placeholder(doc, '图3-7 Agent执行状态图')

    add_body_paragraph(doc, 'Agent辅助写作的时序交互过程为：用户在前端点击"开始写作"按钮，前端调用POST /api/papers/{id}/agent/start接口。后端Orchestrator（编排管理器）按顺序依次调度各个Agent：首先通过CrewAI的kickoff()方法在asyncio线程池中执行文献解析Agent，执行过程中Agent调用DeepSeek API时触发了base.py中的litellm.completion拦截器，拦截器将每个token通过sse_manager.emit()广播到前端。前端通过EventSource监听SSE事件流，实时接收事件类型包括agent_start（Agent开始）、agent_stream（逐token流式输出）、agent_stream_end（流式结束）、agent_complete（Agent完成）/agent_error（Agent异常）和pipeline_complete（全流程完成）。每个Agent完成后，后端将执行结果存入tasks表，前端等待用户审阅反馈。Agent辅助写作时序图如图3-8所示。')

    add_image_placeholder(doc, '图3-8 Agent辅助写作时序图')

    # 3.2.4
    add_heading_para(doc, '3.2.4 文献库管理模块', level=2)

    add_body_paragraph(doc, '文献库管理模块为用户提供个人学术文献的集中管理功能。在文献库页面，用户可以查看自己已添加的所有文献列表，支持按标题和作者进行搜索查询。每条文献展示标题、作者、出处、摘要等基本信息，并提供查看详情和删除操作。用户点击"添加文献"按钮时，弹出添加文献表单，需要填写文献标题、作者、出处、摘要、关键词等信息。用户也可以通过上传文件（支持PDF、Word、TXT格式）自动解析文献内容，上传的文件通过后端文件处理逻辑提取全文文本并存入full_text字段。')

    add_body_paragraph(doc, '用户提交文献后，后端调用POST /api/user/references接口将文献信息存入user_references表。用户可以在文献库页面查看所有已添加的文献，在创建论文时可以从文献库中勾选参考文献进行关联。文献库操作流程图如图3-9所示。')

    add_image_placeholder(doc, '图3-9 文献库操作流程图')

    # 3.2.5
    add_heading_para(doc, '3.2.5 个人中心模块', level=2)

    add_body_paragraph(doc, '个人中心模块为用户提供个人信息管理和密码修改功能。用户进入个人中心页面后，可以查看已登录用户的基本信息，包括用户名、邮箱、头像等。用户可以上传头像图片，修改昵称、邮箱等个人信息。用户点击"修改密码"选项时，系统弹出修改密码表单，用户需要输入旧密码和新密码。前端对两次输入的新密码进行一致性校验，校验通过后调用POST /api/user/change-password接口。后端验证旧密码是否正确（使用bcrypt比对），验证通过后使用bcrypt对新密码进行哈希加密后更新数据库。修改密码成功后，前端提示"密码修改成功"。个人中心操作流程图如图3-10所示。')

    add_image_placeholder(doc, '图3-10 个人中心操作流程图')

    # 3.2.6
    add_heading_para(doc, '3.2.6 管理后台模块', level=2)

    add_body_paragraph(doc, '管理后台模块为管理员角色提供系统管理功能，包括用户管理、论文管理和系统日志查看。管理员登录后可以通过管理后台查看所有注册用户的列表，了解用户注册情况，必要时可管理用户状态。在论文管理页面，管理员可以查看系统所有论文列表，按状态（草稿、执行中、已完成等）筛选论文，了解系统使用情况，必要时可删除异常论文。在系统日志页面，管理员可以查看Agent执行的详细日志，包括每个Agent的输入内容、输出内容、执行耗时、状态等信息，便于监控系统运行状况和排查问题。日志支持按用户、论文、状态等多维度筛选。管理后台还提供系统统计功能，展示用户总数、论文总数等基础统计数据。管理员功能结构如图3-11所示。')

    add_image_placeholder(doc, '图3-11 管理员功能结构图')

    add_body_paragraph(doc, '管理员对用户进行管理的时序交互过程为：管理员登录后进入管理后台用户管理页面，页面加载时调用GET /api/admin/users接口获取所有用户列表。管理员可以点击查看某用户的详细信息，也可执行删除操作。删除用户时调用DELETE /api/admin/users/{id}接口，后端校验操作者身份为管理员后执行删除并返回结果。管理员用户管理时序图如图3-12所示。')

    add_image_placeholder(doc, '图3-12 管理员用户管理时序图')

    # 说明段落
    add_heading_para(doc, '3.2.7 系统接口设计', level=2)

    add_body_paragraph(doc, '本系统后端API采用RESTful风格设计，所有接口路径以/api/开头，按功能模块划分路由。接口返回统一格式的JSON响应，包含success（状态标识）、data（数据负载）和error（错误信息）三个字段。需要身份验证的接口通过JWT Bearer Token进行鉴权，前端在请求头中携带Authorization: Bearer <token>。API路由模块汇总如表3-1所示。')

    add_table_caption(doc, '表3-1 API路由模块汇总表')

    # 创建表格
    table = doc.add_table(rows=8, cols=3, style='Table Grid')
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    headers = ['路由前缀', '模块', '主要功能']
    data = [
        ['/api/auth', '认证管理', '用户注册、登录、获取当前用户信息'],
        ['/api/user', '用户管理', '用户资料查看与更新、密码修改'],
        ['/api/papers', '论文管理', '论文CRUD、论文导出Word'],
        ['/api/papers/{id}/agent', 'Agent执行', '同步/异步执行Agent、SSE流式输出、用户反馈'],
        ['/api/user/references', '文献管理', '文献CRUD、文件上传解析'],
        ['/api/admin', '管理后台', '用户管理、论文管理、系统日志查看'],
    ]

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, font_cn='宋体', font_en='Times New Roman', size=Pt(10.5), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            set_run_font(run, font_cn='宋体', font_en='Times New Roman', size=Pt(10.5))

    # 保存
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'PaperCraft_第3章_系统设计.docx')
    doc.save(output_path)
    print(f'文档已保存至: {output_path}')


if __name__ == '__main__':
    create_chapter3()
