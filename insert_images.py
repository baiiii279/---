#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
批量将SVG渲染为PNG并插入第3章Word文档
"""
import os, sys, glob, re
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
SVG_DIR = os.path.join(BASE, 'diagrams_bw')
PNG_DIR = os.path.join(BASE, 'diagrams_png')
DOCX_SRC = os.path.join(BASE, 'PaperCraft_第3章_系统设计.docx')
DOCX_OUT = os.path.join(BASE, 'PaperCraft_第3章_系统设计_含图.docx')
DOCX_OUT_TMP = DOCX_OUT + '.tmp'  # fallback if locked
os.makedirs(PNG_DIR, exist_ok=True)

# step 1: SVG -> PNG
print('=== Step 1: SVG to PNG ===')
svgs = sorted(glob.glob(os.path.join(SVG_DIR, '图3-*.svg')))
print(f'Found {len(svgs)} SVGs')

def svg_to_png(svg_path, png_path):
    with open(svg_path, 'r', encoding='utf-8') as f:
        svg = f.read()
    # extract viewBox for sizing
    m = re.search(r'viewBox="0 0 (\d+) (\d+)"', svg)
    if m:
        vw, vh = int(m.group(1)), int(m.group(2))
    else:
        vw, vh = 600, 400
    # render at 3x for high DPI
    w, h = vw * 3, vh * 3
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(viewport={'width': w, 'height': h}, device_scale_factor=1)
        html = f'<html><body style="margin:0;background:white">{svg}</body></html>'
        page.set_content(html, wait_until='networkidle')
        page.screenshot(path=png_path, full_page=False)
        browser.close()
    size_kb = os.path.getsize(png_path) // 1024
    print(f'  [PNG] {os.path.basename(png_path)} ({vw}x{vh} -> {w}x{h}, {size_kb}KB)')

for svg_path in svgs:
    basename = os.path.basename(svg_path).replace('.svg', '.png')
    png_path = os.path.join(PNG_DIR, basename)
    if os.path.exists(png_path):
        print(f'  [skip] {basename} (exists)')
        continue
    svg_to_png(svg_path, png_path)

# step 2: insert into Word
print('\n=== Step 2: Insert into Word ===')

doc = Document(DOCX_SRC)

# Build image map
img_map = {}
for f in os.listdir(PNG_DIR):
    if f.endswith('.png'):
        m = re.search(r'(图3-\d+)', f)
        if m:
            img_map[m.group(1)] = os.path.join(PNG_DIR, f)

print(f'Image map: {list(img_map.keys())}')

# Find placeholder paragraphs
to_replace = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith('[此处插入') and '图3-' in text:
        m = re.search(r'(图3-\d+)', text)
        if m and m.group(1) in img_map:
            to_replace.append((i, para, m.group(1)))

print(f'Found {len(to_replace)} placeholders')

# Replace placeholder paragraphs
# We need to insert new paragraphs AFTER the placeholder, then remove the placeholder.
# This avoids all style inheritance issues.
for idx, para, fig_id in reversed(to_replace):
    png_path = img_map[fig_id]

    # Get image dimensions
    from PIL import Image
    with Image.open(png_path) as img:
        pw, ph = img.size

    # Calculate Word dimensions (12cm wide max)
    TW = 12.0
    TH = TW * ph / pw
    if TH > 16.0:
        TH = 16.0
        TW = TH * pw / ph

    # Create a new paragraph after the placeholder using python-docx
    # This properly embeds images with all relationships
    # We use the underlying XML to insert at the right position
    parent = para._element.getparent()
    old_p = para._element

    # Create new paragraph through python-docx for proper image handling
    # First clean the old paragraph (remove all runs and keep as container)
    para.clear()

    # Override line spacing at XML level
    pPr = para._element.get_or_add_pPr()
    # Remove any existing spacing
    for child in list(pPr):
        if child.tag == qn('w:spacing'):
            pPr.remove(child)
    # Set single-line auto spacing
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:line'), '240')
    sp.set(qn('w:lineRule'), 'auto')
    sp.set(qn('w:before'), '60')
    sp.set(qn('w:after'), '60')
    pPr.append(sp)
    # Center alignment
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'center')

    # Add picture directly - this handles all the relationship and blob management
    run = para.add_run()
    run.add_picture(png_path, width=Cm(TW))

    print(f'  [插入] {fig_id} -> {TW:.1f}cm x {TH:.1f}cm')

# Save (try direct, fallback to temp + copy)
try:
    doc.save(DOCX_OUT)
    print(f'\nDone! Saved to: {DOCX_OUT}')
except PermissionError:
    import subprocess
    tmp = DOCX_OUT_TMP
    doc.save(tmp)
    print(f'\nSaved to temp: {tmp}')
    subprocess.run(['powershell', '-c',
        f'Copy-Item "{tmp}" "{DOCX_OUT}" -Force; Remove-Item "{tmp}" -Force'],
        capture_output=True)
    print(f'Copied to: {DOCX_OUT}')
