#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成 PaperCraft 课程设计论文 第4章 数据库设计 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_run_font(run, font_cn='宋体', font_en='Times New Roman', size=Pt(12), bold=None):
    run.font.size = size
    if bold is not None:
        run.font.bold = bold
    run.font.name = font_en
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn('w:eastAsia'), font_cn)


def set_paragraph_spacing(paragraph, line_spacing=20):
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(line_spacing * 20))
    spacing.set(qn('w:lineRule'), 'exact')


def add_body(doc, text, sz=Pt(12)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=sz)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(24)
    set_paragraph_spacing(p)
    return p


def add_heading_l1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_cn='黑体', size=Pt(14), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p)
    return p


def add_heading_l2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_cn='黑体', size=Pt(12), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p)
    return p


def add_heading_l3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_cn='黑体', size=Pt(12), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(24)
    set_paragraph_spacing(p)
    return p


def add_table_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=Pt(10.5), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p)
    return p


def add_fig_caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=Pt(10.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p)
    return p


def add_img_placeholder(doc, label):
    p = doc.add_paragraph()
    run = p.add_run(f'[此处插入{label}]')
    set_run_font(run, size=Pt(10.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.color.rgb = RGBColor(128, 128, 128)
    set_paragraph_spacing(p)
    return p


def make_table(doc, headers, rows, col_widths=None):
    """添加表格，表头加粗，内容五号字"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=Pt(10.5), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=Pt(10.5))
    return table


def create_ch4():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')

    # ========== 第4章 ==========
    add_heading_l1(doc, '第4章 数据库设计')

    add_body(doc, '数据库设计是系统设计的核心环节之一，其目标是建立一个高效、安全、一致的数据存储结构，以支撑系统各功能模块的正常运行。本系统采用MySQL关系型数据库，通过SQLAlchemy ORM框架进行对象关系映射，实现了业务逻辑与数据访问的分离。本章将从概念设计和逻辑设计两个层面，详细阐述系统的数据库设计方案。本章内容对应系统中的6张数据表。')

    # ========== 4.1 ==========
    add_heading_l2(doc, '4.1 数据库概念设计')

    add_body(doc, '数据库概念设计是对系统业务对象及其关系进行抽象建模的过程，主要用于描述系统中各类实体、实体属性以及实体之间的联系。本系统围绕普通用户和管理员两类角色展开，核心业务实体包括用户、论文、文献、任务和Agent日志。实体之间的联系为：一个用户可以创建多篇论文；一篇论文可以关联多篇参考文献，一篇文献也可以被多篇论文引用；一篇论文包含多个Agent执行任务；一个Agent任务包含多条执行日志。')

    # 4.1.1
    add_heading_l3(doc, '4.1.1 实体和属性的定义')

    add_body(doc, '本系统涉及的核心实体共6个，分别为用户（User）、论文（Paper）、用户文献（UserReference）、论文-文献关联（PaperReference）、任务（Task）和Agent日志（AgentLog）。各实体的属性定义如下：')

    entities = [
        ('用户（User）', '用户ID、用户名、邮箱、密码哈希、角色（user/admin）、头像URL、创建时间、更新时间'),
        ('论文（Paper）', '论文ID、用户ID（外键）、标题、主题、模板类型（course/journal）、状态（draft→complete七阶段）、大纲JSON、正文Markdown、创建时间、更新时间'),
        ('用户文献（UserReference）', '文献ID、用户ID（外键）、标题、作者、出处、摘要、全文内容、URL、关键词、创建时间'),
        ('论文-文献关联（PaperReference）', '关联ID、论文ID（外键）、文献ID（外键）'),
        ('任务（Task）', '任务ID、论文ID（外键）、Agent类型（parse/outline/write/polish/cite_check）、状态（pending/running/success/failed）、输入数据JSON、输出数据JSON、用户反馈状态、反馈意见、开始时间、完成时间'),
        ('Agent日志（AgentLog）', '日志ID（BIGINT）、任务ID（外键）、步骤名、消息内容、级别（info/warn/error）、创建时间'),
    ]

    for name, attrs in entities:
        add_body(doc, f'{name}，包含属性：{attrs}。', sz=Pt(12))

    add_body(doc, '在实体关系中，一个用户可以创建多篇论文，因此用户与论文之间是一对多（1:N）关系。一篇论文可以引用多篇文献，一篇文献也可以被多篇论文引用，因此论文与文献之间是多对多（N:M）关系，通过论文-文献关联表实现。一篇论文在写作过程中会产生多个Agent任务，因此论文与任务之间是一对多（1:N）关系。一个Agent任务在执行过程中会产生多条日志记录，因此任务与日志之间也是一对多（1:N）关系。')

    # 4.1.2
    add_heading_l3(doc, '4.1.2 E-R模式设计')

    add_body(doc, 'E-R模型（Entity-Relationship Model）为建立数据库实体关系提供了直观的图形化描述方法。本系统的全局E-R图如图4-1所示，展示了各实体之间的关联关系。')

    add_img_placeholder(doc, '图4-1 系统全局E-R图')

    add_body(doc, '用户作为系统的核心实体，与论文和文献两个实体直接关联。用户实体包含用户ID（主键）、用户名、邮箱、密码哈希、角色、头像URL、创建时间和更新时间等属性。用户可以通过系统创建多篇论文，也可以管理自己的文献库。用户实体的属性图如图4-2所示。')

    add_img_placeholder(doc, '图4-2 用户实体属性图')

    add_body(doc, '论文实体包含论文ID（主键）、用户ID（外键）、标题、主题、模板类型、状态、大纲和正文内容等属性。论文状态从草稿（draft）开始，经过解析中（parsing）、大纲生成中（outlining）、内容撰写中（writing）、润色中（polishing）、检查中（checking），最终变为已完成（complete），共七个状态。论文实体的属性图如图4-3所示。')

    add_img_placeholder(doc, '图4-3 论文实体属性图')

    add_body(doc, '用户文献实体包含文献ID（主键）、用户ID（外键）、标题、作者、出处、摘要、全文内容、URL、关键词和创建时间等属性。文献的全文内容通过文件上传解析获得，可支持PDF、Word和TXT格式的自动文本提取。用户文献实体属性图如图4-4所示。')

    add_img_placeholder(doc, '图4-4 用户文献实体属性图')

    add_body(doc, '任务实体是Agent执行过程的记录载体，包含任务ID（主键）、论文ID（外键）、Agent类型、状态、输入输出数据、用户反馈状态和反馈意见等属性。Agent类型包括parse（文献解析）、outline（大纲生成）、write（内容撰写）、polish（润色）和cite_check（引用检查）五种。用户可以对每个Agent的执行结果做出批准（approve）、驳回（reject）或编辑（edit）的反馈。任务实体属性图如图4-5所示。')

    add_img_placeholder(doc, '图4-5 任务实体属性图')

    # ========== 4.2 ==========
    add_heading_l2(doc, '4.2 数据库逻辑设计')

    add_body(doc, '数据库逻辑设计是将概念设计阶段得到的E-R模型转换为关系数据库管理系统所支持的关系模型的过程。本系统采用MySQL关系型数据库，通过SQLAlchemy ORM框架进行数据持久化操作，所有数据表均遵循第三范式（3NF）进行设计，避免了数据冗余和更新异常。系统共包含6张数据表，总表设计如表4-1所示。')

    add_table_title(doc, '表4-1 数据总表')

    make_table(doc,
        ['表名称', '实际表名', '说明'],
        [
            ('用户表', 'users', '存储用户和管理员的账号信息与角色'),
            ('论文表', 'papers', '存储论文的元数据、内容和状态'),
            ('用户文献表', 'user_references', '存储用户上传的参考文献信息'),
            ('论文-文献关联表', 'paper_references', '论文与文献的多对多关联'),
            ('任务表', 'tasks', '存储Agent执行的输入输出和用户反馈'),
            ('Agent日志表', 'agent_logs', '存储Agent执行的详细日志记录'),
        ]
    )

    # ---- users ----
    add_body(doc, '用户表users存储所有用户和管理员的账号信息，包括用户ID、用户名、邮箱、密码哈希、角色、头像等字段。角色字段使用枚举类型（user/admin），默认为普通用户。密码使用SHA-256加盐哈希算法存储，不保存明文密码。用户表设计如表4-2所示。')

    add_table_title(doc, '表4-2 用户表users')

    make_table(doc,
        ['序号', '字段名称', '数据类型', '长度', '允许空', '备注'],
        [
            ('1', 'id', 'INT', '11', '否', '用户ID（主键，自增）'),
            ('2', 'username', 'VARCHAR', '50', '否', '用户名（唯一）'),
            ('3', 'email', 'VARCHAR', '100', '是', '邮箱（唯一）'),
            ('4', 'password_hash', 'VARCHAR', '255', '否', '密码哈希值'),
            ('5', 'role', 'ENUM', '-', '否', '角色：user/admin（默认user）'),
            ('6', 'avatar', 'VARCHAR', '255', '是', '头像URL'),
            ('7', 'created_at', 'DATETIME', '-', '否', '创建时间'),
            ('8', 'updated_at', 'DATETIME', '-', '否', '更新时间'),
        ]
    )

    # ---- papers ----
    add_body(doc, '论文表papers是系统的核心业务表，记录用户提交的每一篇论文。论文状态严谨遵循七阶段流转：draft→parsing→outlining→writing→polishing→checking→complete。状态流转通过Orchestrator编排器自动管理。论文表设计如表4-3所示。')

    add_table_title(doc, '表4-3 论文表papers')

    make_table(doc,
        ['序号', '字段名称', '数据类型', '长度', '允许空', '备注'],
        [
            ('1', 'id', 'INT', '11', '否', '论文ID（主键，自增）'),
            ('2', 'user_id', 'INT', '11', '否', '用户ID（外键→users.id）'),
            ('3', 'title', 'VARCHAR', '200', '是', '论文标题'),
            ('4', 'topic', 'VARCHAR', '500', '否', '论文主题'),
            ('5', 'template', 'ENUM', '-', '否', '模板：course/journal'),
            ('6', 'status', 'ENUM', '-', '否', '状态：七阶段流转（默认draft）'),
            ('7', 'outline', 'TEXT', '-', '是', 'JSON格式论文大纲'),
            ('8', 'content', 'TEXT/LONGTEXT', '-', '是', 'Markdown格式正文'),
            ('9', 'created_at', 'DATETIME', '-', '否', '创建时间'),
            ('10', 'updated_at', 'DATETIME', '-', '否', '更新时间'),
        ]
    )

    # ---- user_references ----
    add_body(doc, '用户文献表user_references存储用户添加或上传的学术参考文献。用户可以通过手动填写或文件上传（PDF/Word/TXT）的方式添加文献，上传的文件经后端解析后将全文文本存储到full_text字段，供文献解析Agent使用。文献表设计如表4-4所示。')

    add_table_title(doc, '表4-4 用户文献表user_references')

    make_table(doc,
        ['序号', '字段名称', '数据类型', '长度', '允许空', '备注'],
        [
            ('1', 'id', 'INT', '11', '否', '文献ID（主键，自增）'),
            ('2', 'user_id', 'INT', '11', '否', '用户ID（外键→users.id）'),
            ('3', 'title', 'VARCHAR', '300', '否', '文献标题'),
            ('4', 'authors', 'VARCHAR', '500', '是', '作者'),
            ('5', 'source', 'VARCHAR', '500', '是', '出处（期刊/会议等）'),
            ('6', 'abstract', 'TEXT', '-', '是', '摘要'),
            ('7', 'full_text', 'TEXT/LONGTEXT', '-', '是', '全文内容（文件上传后提取）'),
            ('8', 'url', 'VARCHAR', '500', '是', '文献链接'),
            ('9', 'keywords', 'VARCHAR', '300', '是', '关键词'),
            ('10', 'created_at', 'DATETIME', '-', '否', '创建时间'),
        ]
    )

    # ---- paper_references ----
    add_body(doc, '论文-文献关联表paper_references用于实现论文与文献之间的多对多关系。每篇论文可以引用多篇文献，每篇文献也可以被多篇论文引用。当用户创建论文时，系统将选中的文献关联写入该表。当论文或文献被删除时，关联数据通过级联删除（CASCADE）自动清理。关联表设计如表4-5所示。')

    add_table_title(doc, '表4-5 论文-文献关联表paper_references')

    make_table(doc,
        ['序号', '字段名称', '数据类型', '长度', '允许空', '备注'],
        [
            ('1', 'id', 'INT', '11', '否', '关联ID（主键，自增）'),
            ('2', 'paper_id', 'INT', '11', '否', '论文ID（外键→papers.id，级联删除）'),
            ('3', 'reference_id', 'INT', '11', '否', '文献ID（外键→user_references.id，级联删除）'),
        ]
    )

    # ---- tasks ----
    add_body(doc, '任务表tasks记录每个Agent的执行情况和用户反馈，是实现Human-in-the-loop协作模式的关键数据载体。每个Agent执行完成后，系统将执行结果写入output_data字段，并将用户反馈状态置为pending（待处理），等待用户审阅。用户选择批准（approve）后进入下一Agent；选择驳回（reject）时，Agent根据反馈意见重新生成。任务表设计如表4-6所示。')

    add_table_title(doc, '表4-6 任务表tasks')

    make_table(doc,
        ['序号', '字段名称', '数据类型', '长度', '允许空', '备注'],
        [
            ('1', 'id', 'INT', '11', '否', '任务ID（主键，自增）'),
            ('2', 'paper_id', 'INT', '11', '否', '论文ID（外键→papers.id）'),
            ('3', 'agent_type', 'ENUM', '-', '否', 'Agent类型：parse/outline/write/polish/cite_check'),
            ('4', 'status', 'ENUM', '-', '否', '状态：pending/running/success/failed'),
            ('5', 'input_data', 'TEXT', '-', '是', 'Agent输入JSON'),
            ('6', 'output_data', 'TEXT/LONGTEXT', '-', '是', 'Agent输出内容'),
            ('7', 'user_feedback', 'ENUM', '-', '否', '用户反馈：pending/approve/reject/edit'),
            ('8', 'feedback_comment', 'TEXT', '-', '是', '驳回/编辑的修改意见'),
            ('9', 'started_at', 'DATETIME', '-', '是', '任务开始时间'),
            ('10', 'finished_at', 'DATETIME', '-', '是', '任务完成时间'),
        ]
    )

    # ---- agent_logs ----
    add_body(doc, 'Agent日志表agent_logs记录Agent执行过程中产生的每一条日志，用于系统监控和问题排查。日志的级别分为info（信息）、warn（警告）和error（错误）三种。管理员可以在管理后台按级别、用户和论文进行日志筛选。日志表设计如表4-7所示。')

    add_table_title(doc, '表4-7 Agent日志表agent_logs')

    make_table(doc,
        ['序号', '字段名称', '数据类型', '长度', '允许空', '备注'],
        [
            ('1', 'id', 'BIGINT', '20', '否', '日志ID（主键，自增）'),
            ('2', 'task_id', 'INT', '11', '否', '任务ID（外键→tasks.id）'),
            ('3', 'step', 'VARCHAR', '50', '否', '执行步骤名'),
            ('4', 'message', 'TEXT', '-', '否', '日志消息内容'),
            ('5', 'level', 'ENUM', '-', '否', '级别：info/warn/error（默认info）'),
            ('6', 'created_at', 'DATETIME', '-', '否', '创建时间'),
        ]
    )

    add_body(doc, '综上所述，本系统的6张数据表通过外键关联形成了完整的数据关系网络。用户表与其他表之间通过用户ID关联；论文表与任务表之间通过论文ID关联；论文表与文献表之间通过关联表建立多对多关系；任务表与日志表之间通过任务ID关联。所有表的创建时间字段均使用数据库的自动当前时间戳（server_default=func.now()），数据完整性通过外键约束和级联删除机制保障。')

    # 保存
    output = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'PaperCraft_第4章_数据库设计.docx')
    doc.save(output)
    print(f'文档已保存: {output}')


if __name__ == '__main__':
    create_ch4()
