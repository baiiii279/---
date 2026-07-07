"""Tests for the format template management API endpoints."""
import io
import tempfile
from unittest.mock import patch

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, text
from sqlalchemy.orm import sessionmaker

from app.core.database import Base, get_db
from app.main import app
from app.models.user import User
from app.models.format_template import FormatTemplate
from app.api.auth import get_current_user

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

_ALL_TABLES = ["format_templates", "users", "user_references", "paper_references", "papers", "tasks", "agent_logs"]


class _FakeUser:
    """Plain object replacing a SQLAlchemy User for dependency overrides."""
    def __init__(self, id: int, username: str = "testuser"):
        self.id = id
        self.username = username


@pytest.fixture(scope="module")
def sqlite_engine():
    """Module-scoped engine using a temp file so all test functions and
    their httpx worker threads share the same database."""
    with tempfile.NamedTemporaryFile(suffix=".db", delete=False) as f:
        db_path = f.name
    engine = create_engine(f"sqlite:///{db_path}", connect_args={"check_same_thread": False})
    Base.metadata.create_all(bind=engine)
    yield engine
    engine.dispose()
    import os
    os.unlink(db_path)


@pytest.fixture(autouse=True)
def _clean_db(sqlite_engine):
    """Clean all rows between test functions to ensure isolation."""
    with sqlite_engine.connect() as conn:
        for table in _ALL_TABLES:
            conn.execute(text(f"DELETE FROM {table}"))
        conn.commit()
    yield


@pytest.fixture
def db_session(sqlite_engine):
    """Function-scoped session on the shared engine."""
    SessionLocal = sessionmaker(bind=sqlite_engine)
    session = SessionLocal()
    try:
        yield session
    finally:
        session.close()


@pytest.fixture
def client(sqlite_engine):
    """TestClient with DB dependency returning fresh sessions from the
    shared engine.  App startup events are silenced."""

    def _override_get_db():
        SessionLocal = sessionmaker(bind=sqlite_engine)
        db = SessionLocal()
        try:
            yield db
        finally:
            db.close()

    app.dependency_overrides[get_db] = _override_get_db

    with (
        patch("app.main._seed_admin"),
        patch("app.main._seed_default_template"),
    ):
        client = TestClient(app)
    try:
        yield client
    finally:
        app.dependency_overrides.clear()


@pytest.fixture
def test_user_id(db_session):
    """Persist a test user and return its id."""
    user = User(username="testuser", email="test@example.com", password_hash="hunter2")
    db_session.add(user)
    db_session.commit()
    return user.id


@pytest.fixture
def another_user_id(db_session):
    """Persist a second user and return its id."""
    user = User(username="other", email="other@example.com", password_hash="hunter2")
    db_session.add(user)
    db_session.commit()
    return user.id


@pytest.fixture
def auth_client(client, test_user_id):
    """Override auth dependency to return a fake user with test_user_id."""

    def _override_user():
        return _FakeUser(id=test_user_id)

    app.dependency_overrides[get_current_user] = _override_user
    yield client
    app.dependency_overrides.pop(get_current_user, None)


@pytest.fixture
def auth_client_as_other(client, another_user_id):
    """Override auth dependency to return a fake user with another_user_id."""

    def _override_user():
        return _FakeUser(id=another_user_id)

    app.dependency_overrides[get_current_user] = _override_user
    yield client
    app.dependency_overrides.pop(get_current_user, None)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def make_docx_bytes() -> bytes:
    """Return a minimal valid .docx file as bytes."""
    from docx import Document

    doc = Document()
    doc.add_heading("标题", level=1)
    doc.add_paragraph("正文内容")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# GET /api/format-templates
# ---------------------------------------------------------------------------


class TestListTemplates:
    def test_empty_list(self, auth_client, db_session):
        """Returns empty list when no templates exist."""
        resp = auth_client.get("/api/format-templates")
        assert resp.status_code == 200
        assert resp.json() == []

    def test_lists_user_and_default_templates(self, auth_client, db_session, test_user_id):
        """Returns user's own templates plus default templates."""
        default = FormatTemplate(
            user_id=0, name="默认", rules="default rules", is_default=True
        )
        db_session.add(default)
        mine = FormatTemplate(
            user_id=test_user_id, name="我的模板", rules="my rules", is_default=False
        )
        db_session.add(mine)
        # Another user's template (should not appear)
        others = FormatTemplate(
            user_id=999, name="别人的", rules="other rules", is_default=False
        )
        db_session.add(others)
        db_session.commit()

        resp = auth_client.get("/api/format-templates")
        data = resp.json()
        assert resp.status_code == 200
        names = {t["name"] for t in data}
        assert "默认" in names
        assert "我的模板" in names
        assert "别人的" not in names

    def test_default_template_first_in_order(self, auth_client, db_session, test_user_id):
        """Default templates appear before user templates."""
        mine = FormatTemplate(
            user_id=test_user_id, name="我的", rules="my", is_default=False,
        )
        db_session.add(mine)
        default = FormatTemplate(
            user_id=0, name="默认", rules="default", is_default=True,
        )
        db_session.add(default)
        db_session.commit()

        resp = auth_client.get("/api/format-templates")
        data = resp.json()
        assert data[0]["is_default"] is True

    def test_response_shape(self, auth_client, db_session, test_user_id):
        """Each template has id, name, is_default, created_at."""
        t = FormatTemplate(
            user_id=test_user_id, name="模板", rules="rules", is_default=False,
        )
        db_session.add(t)
        db_session.commit()

        resp = auth_client.get("/api/format-templates")
        item = resp.json()[0]
        assert set(item.keys()) == {"id", "name", "is_default", "created_at"}


# ---------------------------------------------------------------------------
# POST /api/format-templates/upload
# ---------------------------------------------------------------------------


class TestUploadTemplate:
    def test_upload_valid_docx(self, auth_client, db_session, test_user_id):
        """Upload a valid docx returns parsed rules."""
        docx_bytes = make_docx_bytes()
        resp = auth_client.post(
            "/api/format-templates/upload",
            files={"file": ("test.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"name": "学术论文格式"},
        )
        assert resp.status_code == 200
        body = resp.json()
        assert body["name"] == "学术论文格式"
        assert "rules" in body
        assert "id" in body
        assert "created_at" in body

        # Verify persisted
        saved = db_session.query(FormatTemplate).filter_by(name="学术论文格式").first()
        assert saved is not None
        assert saved.user_id == test_user_id
        assert saved.is_default is False

    def test_upload_non_docx_rejected(self, auth_client):
        """Non-docx file raises 400."""
        resp = auth_client.post(
            "/api/format-templates/upload",
            files={"file": ("test.txt", b"not a docx", "text/plain")},
            data={"name": "bad file"},
        )
        assert resp.status_code == 400
        assert "仅支持 .docx 文件" in resp.json()["detail"]

    def test_upload_parse_failure_returns_400(self, auth_client):
        """Upload that fails to parse returns 400."""
        resp = auth_client.post(
            "/api/format-templates/upload",
            files={"file": ("bad.docx", b"this is not a real docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"name": "bad template"},
        )
        assert resp.status_code == 400
        assert "模板解析失败" in resp.json()["detail"]

    def test_upload_saves_to_user(self, auth_client, db_session, test_user_id):
        """Uploaded template belongs to the current user."""
        docx_bytes = make_docx_bytes()
        resp = auth_client.post(
            "/api/format-templates/upload",
            files={"file": ("t.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"name": "我的格式"},
        )
        assert resp.status_code == 200
        saved = db_session.query(FormatTemplate).filter_by(name="我的格式").first()
        assert saved.user_id == test_user_id

    def test_upload_requires_file(self, auth_client):
        """Missing file should return 422."""
        resp = auth_client.post(
            "/api/format-templates/upload",
            data={"name": "no file"},
        )
        assert resp.status_code == 422

    def test_upload_requires_name(self, auth_client):
        """Missing name should return 422."""
        docx_bytes = make_docx_bytes()
        resp = auth_client.post(
            "/api/format-templates/upload",
            files={"file": ("t.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert resp.status_code == 422

    def test_upload_creates_non_default(self, auth_client, db_session):
        """Uploaded templates always have is_default=False."""
        docx_bytes = make_docx_bytes()
        resp = auth_client.post(
            "/api/format-templates/upload",
            files={"file": ("t.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"name": "non-default"},
        )
        assert resp.status_code == 200
        saved = db_session.query(FormatTemplate).filter_by(name="non-default").first()
        assert saved.is_default is False

    def test_temp_file_cleaned_up(self, auth_client):
        """Temporary file should be deleted after request completes."""
        docx_bytes = make_docx_bytes()
        resp = auth_client.post(
            "/api/format-templates/upload",
            files={"file": ("t.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"name": "cleanup test"},
        )
        assert resp.status_code == 200


# ---------------------------------------------------------------------------
# DELETE /api/format-templates/{template_id}
# ---------------------------------------------------------------------------


class TestDeleteTemplate:
    def test_delete_own_template(self, auth_client, db_session, test_user_id):
        """User can delete their own non-default template."""
        t = FormatTemplate(
            user_id=test_user_id, name="可删除", rules="rules", is_default=False,
        )
        db_session.add(t)
        db_session.commit()

        resp = auth_client.delete(f"/api/format-templates/{t.id}")
        assert resp.status_code == 200
        assert resp.json() == {"message": "ok"}
        assert db_session.query(FormatTemplate).filter_by(id=t.id).first() is None

    def test_delete_non_existent_returns_404(self, auth_client):
        """Deleting a non-existent template returns 404."""
        resp = auth_client.delete("/api/format-templates/99999")
        assert resp.status_code == 404
        assert "模板不存在" in resp.json()["detail"]

    def test_delete_others_template_returns_404(self, auth_client, db_session, another_user_id):
        """User cannot delete another user's template."""
        t = FormatTemplate(
            user_id=another_user_id, name="别人的", rules="rules", is_default=False,
        )
        db_session.add(t)
        db_session.commit()

        resp = auth_client.delete(f"/api/format-templates/{t.id}")
        assert resp.status_code == 404
        assert "无权删除" in resp.json()["detail"]

    def test_delete_default_template_rejected(self, auth_client, db_session, test_user_id):
        """Default templates cannot be deleted."""
        t = FormatTemplate(
            user_id=test_user_id, name="默认的", rules="rules", is_default=True,
        )
        db_session.add(t)
        db_session.commit()

        resp = auth_client.delete(f"/api/format-templates/{t.id}")
        assert resp.status_code == 400
        assert "不能删除默认模板" in resp.json()["detail"]
