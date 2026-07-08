import re
import io
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from app.core.database import get_db
from app.models.paper import Paper
from app.models.reference import PaperReference
from app.models.task import Task
from app.models.agent_log import AgentLog
from app.models.user import User
from app.schemas.paper import CreatePaperRequest, PaperResponse, UpdatePaperRequest
from app.api.auth import get_current_user

router = APIRouter(prefix="/api/papers", tags=["papers"])


@router.post("", response_model=PaperResponse)
def create_paper(req: CreatePaperRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    paper = Paper(user_id=current_user.id, topic=req.topic, template=req.template)
    db.add(paper)
    db.flush()

    for ref_id in req.reference_ids:
        db.add(PaperReference(paper_id=paper.id, reference_id=ref_id))

    db.commit()
    db.refresh(paper)
    return paper


@router.get("", response_model=list[PaperResponse])
def list_papers(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(Paper).filter(Paper.user_id == current_user.id).order_by(Paper.created_at.desc()).all()


@router.get("/{paper_id}", response_model=PaperResponse)
def get_paper(paper_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    paper = db.query(Paper).filter(Paper.id == paper_id, Paper.user_id == current_user.id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="论文不存在")
    return paper


@router.put("/{paper_id}")
def update_paper(paper_id: int, req: UpdatePaperRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    paper = db.query(Paper).filter(Paper.id == paper_id, Paper.user_id == current_user.id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="论文不存在")
    if req.title is not None:
        paper.title = req.title
    if req.outline is not None:
        paper.outline = req.outline
    if req.content is not None:
        paper.content = req.content
    db.commit()
    return {"message": "ok"}


@router.delete("/{paper_id}")
def delete_paper(paper_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    paper = db.query(Paper).filter(Paper.id == paper_id, Paper.user_id == current_user.id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="论文不存在")
    # 先删除关联的任务和日志
    task_ids = [t.id for t in db.query(Task).filter(Task.paper_id == paper_id).all()]
    if task_ids:
        db.query(AgentLog).filter(AgentLog.task_id.in_(task_ids)).delete(synchronize_session=False)
        db.query(Task).filter(Task.paper_id == paper_id).delete(synchronize_session=False)
    # 再删除论文关联的文献引用和论文本身
    db.query(PaperReference).filter(PaperReference.paper_id == paper_id).delete()
    db.delete(paper)
    db.commit()
    return {"message": "ok"}


def _md_to_docx(md_text: str, title: str) -> io.BytesIO:
    """将 Markdown 文本转换为 Word 文档"""
    doc = Document()

    # 设置默认字体（全黑）
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)  # 小四
    font.color.rgb = BLACK
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    lines = md_text.split('\n')
    _current_format = None
    _has_content = False
    i = 0
    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 识别排版指令标记 (来自 FormatAgent)
        format_cmd_match = re.match(r'<!-- format:\s*(\S+)', line)
        if format_cmd_match:
            cmd = format_cmd_match.group(1)
            if cmd == 'page-break':
                # 分页符（只在已有内容后插入，避免开头和连续白页）
                if _has_content:
                    doc.add_page_break()
            elif cmd == 'body-text':
                # 正文模式（默认小四宋体）
                _current_format = 'body'
            elif cmd == 'ref-list':
                # 参考文献列表（缩小字号为五号）
                _current_format = 'ref'
            i += 1
            continue

        # 遇到正文内容
        _has_content = True

        # [摘要] 和 [关键词] 特殊处理 — 居中加粗
        if line.startswith('[摘要]') or line.startswith('[关键词]'):
            text = line.strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font([run], '黑体', Pt(14), bold=True)
            i += 1
            continue

        # H1 标题（论文题目，居中）
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(p.runs, '黑体', Pt(16))
            i += 1
            continue

        # H2 标题（章节标题，左对齐）
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_run_font(p.runs, '黑体', Pt(14))
            i += 1
            continue

        # H3 标题
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            _set_run_font(p.runs, '黑体', Pt(13))
            i += 1
            continue

        # 分隔线 — 跳过，不渲染任何内容
        if line.strip() == '---':
            i += 1
            continue

        # 表格 (检测 | 开头的行)
        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # 过滤分隔行 (如 |---|---|)
            data_rows = [tl for tl in table_lines if not re.match(r'^\|[\s\-:|]+\|$', tl)]
            if data_rows:
                num_cols = len(data_rows[0].split('|')) - 2  # 去掉首尾空
                table = doc.add_table(rows=len(data_rows), cols=max(1, num_cols))
                table.style = 'Table Grid'
                for ri, row_line in enumerate(data_rows):
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    for ci, cell_text in enumerate(cells[:num_cols]):
                        if ri == 0:
                            # Header row - bold
                            cell = table.rows[ri].cells[ci]
                            cell.text = ''
                            run = cell.paragraphs[0].add_run(cell_text)
                            run.bold = True
                            run.font.size = Pt(10)
                            run.font.name = '宋体'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            # 灰色背景
                            shading_elm = cell._element.get_or_add_tcPr()
                            shading = shading_elm.makeelement(qn('w:shd'), {
                                qn('w:val'): 'clear',
                                qn('w:color'): 'auto',
                                qn('w:fill'): 'D9E2F3'
                            })
                            shading_elm.append(shading)
                        else:
                            cell = table.rows[ri].cells[ci]
                            cell.text = ''
                            run = cell.paragraphs[0].add_run(cell_text)
                            run.font.size = Pt(10)
                            run.font.name = '宋体'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            continue

        # 无序列表
        if re.match(r'^[\*\-\+]\s+', line.strip()):
            text = re.sub(r'^[\*\-\+]\s+', '', line.strip())
            p = doc.add_paragraph(text, style='List Bullet')
            _format_paragraph(p)
            if _current_format == 'ref':
                for run in p.runs:
                    run.font.size = Pt(10)
            i += 1
            continue

        # 有序列表
        if re.match(r'^\d+[\.\)]\s', line.strip()):
            text = re.sub(r'^\d+[\.\)]\s+', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
            _format_paragraph(p)
            if _current_format == 'ref':
                for run in p.runs:
                    run.font.size = Pt(10)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_formatted_text(p, line)
        _format_paragraph(p)
        if _current_format == 'ref':
            for run in p.runs:
                run.font.size = Pt(10)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


BLACK = RGBColor(0, 0, 0)


def _set_run_font(runs, font_name: str, font_size: Pt):
    """设置 run 的字体（全黑）"""
    for run in runs:
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = BLACK
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _format_paragraph(p):
    """格式化段落：首行缩进、行距"""
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.74)  # 两个字符
    pf.line_spacing = Pt(20)  # 固定值20磅
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def _add_formatted_text(paragraph, text: str):
    """解析并添加带格式的文本 (粗体、斜体、上标引用)，全部黑色"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.font.color.rgb = BLACK
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sp in sub_parts:
                if sp.startswith('*') and sp.endswith('*') and not sp.startswith('**'):
                    run = paragraph.add_run(sp[1:-1])
                    run.italic = True
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
                    run.font.color.rgb = BLACK
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    ref_parts = re.split(r'(\[\d+(?:,\d+)*\])', sp)
                    for rp in ref_parts:
                        if re.match(r'^\[\d+(?:,\d+)*\]$', rp):
                            run = paragraph.add_run(rp)
                            run.font.superscript = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = BLACK
                        else:
                            run = paragraph.add_run(rp)
                            run.font.name = '宋体'
                            run.font.size = Pt(12)
                            run.font.color.rgb = BLACK
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


@router.get("/{paper_id}/export/docx")
def export_docx(paper_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    """导出论文为 Word 文档"""
    paper = db.query(Paper).filter(Paper.id == paper_id, Paper.user_id == current_user.id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="论文不存在")
    if not paper.content:
        raise HTTPException(status_code=400, detail="论文内容为空，请先运行 Agent 生成内容")

    content = paper.content
    # 清理润色 Agent 的输出产物
    content = re.sub(r'^好的，.*?(?=#{1,3}\s)', '', content.strip(), count=1, flags=re.DOTALL)
    content = re.sub(r'#+\s*\*{0,2}润色后的论文全文\*{0,2}\s*\n*', '', content)
    content = re.sub(r'\n---\n\s*#+\s*\*{0,2}修改说明\*{0,2}.*$', '', content, flags=re.DOTALL)
    content = re.sub(r'\n#+\s+修改说明\s*\n.*$', '', content, flags=re.DOTALL)
    content = re.sub(r'\n通过这些修改，.*$', '', content, flags=re.DOTALL)
    content = re.sub(r'\n{3,}', '\n\n', content)
    content = content.strip()
    if not content:
        content = paper.content

    filename = f"{paper.topic or '论文'}.docx"
    # URL encode Chinese filename
    from urllib.parse import quote
    encoded_filename = quote(filename)

    buf = _md_to_docx(content, paper.topic or '论文')

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )
