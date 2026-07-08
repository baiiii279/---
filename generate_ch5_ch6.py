#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成第5章 系统实现 和 第6章 系统测试"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS = os.path.join(BASE, 'screenshots')

def sf(run, cn='宋体', en='Times New Roman', sz=Pt(12), bold=None):
    run.font.size = sz
    if bold is not None: run.font.bold = bold
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn('w:eastAsia'), cn)

def sp(para, lbs=20):
    pPr = para._element.get_or_add_pPr()
    sp_el = pPr.find(qn('w:spacing'))
    if sp_el is None:
        sp_el = OxmlElement('w:spacing')
        pPr.append(sp_el)
    sp_el.set(qn('w:line'), str(lbs * 20))
    sp_el.set(qn('w:lineRule'), 'exact')

def body(doc, text, sz=Pt(12), indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sf(run, sz=sz)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Pt(24)
    sp(p)
    return p

def h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sf(run, cn='黑体', sz=Pt(14), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sf(run, cn='黑体', sz=Pt(12), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(p)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sf(run, cn='黑体', sz=Pt(12), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(24)
    sp(p)
    return p

def cap(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sf(run, sz=Pt(10.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p)
    return p

def img(doc, name, w=12):
    path = os.path.join(SCREENSHOTS, name)
    if not os.path.exists(path):
        body(doc, f'[截图 {name} 未找到]')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(w))
    pPr = p._element.get_or_add_pPr()
    sp_el = OxmlElement('w:spacing')
    sp_el.set(qn('w:line'), '240')
    sp_el.set(qn('w:lineRule'), 'auto')
    sp_el.set(qn('w:before'), '60')
    sp_el.set(qn('w:after'), '60')
    pPr.append(sp_el)

def tbl(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        sf(run, sz=Pt(10.5), bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            sf(run, sz=Pt(10.5))
    return table

def ttitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sf(run, sz=Pt(10.5), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p)
    return p


# ==================== 第5章 ====================
def create_ch5():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')

    h1(doc, '第5章 系统实现')
    body(doc, '本章将对系统各功能模块的实现进行详细展示，通过系统截图与功能描述相结合的方式，介绍PaperCraft多智能体论文写作助手的核心功能实现。系统采用前后端分离架构，前端使用React 18+TypeScript构建单页应用，后端使用FastAPI框架提供RESTful API服务，CrewAI框架负责编排6个AI Agent协同完成论文写作任务。')

    # 5.1
    h2(doc, '5.1 登录与注册功能')
    body(doc, '登录与注册功能是系统的基础入口模块。用户可以在登录页面输入用户名和密码进行登录，系统验证用户名和密码的正确性后返回JWT令牌用于后续请求的身份认证。登录页面如图5-1所示。')
    img(doc, '09_登录页.png', 10)
    cap(doc, '图5-1 登录页面')
    body(doc, '如果用户没有账号，可以点击"注册"链接进入注册页面。注册时需要填写用户名（至少2个字符）、密码（至少6个字符）和选填的邮箱。系统在前端对输入格式进行前端校验，后端对用户名唯一性进行检查。注册页面如图5-2所示。')
    img(doc, '10_注册页.png', 10)
    cap(doc, '图5-2 注册页面')

    # 5.2
    h2(doc, '5.2 首页')
    body(doc, '首页是用户登录后进入的第一个页面，展示了系统的整体介绍和功能入口。页面顶部为深色导航栏，包含"我的论文"、"文献库"、"个人中心"等链接，管理员用户还可看到"管理后台"入口。首页主体区域展示了PaperCraft的6项核心功能卡片（文献解析、大纲生成、内容撰写、润色优化、引用检查、格式排版），底部提供了"三步完成论文"的快速指引。首页页面如图5-3所示。')
    img(doc, '01_首页.png', 12)
    cap(doc, '图5-3 首页')

    # 5.3
    h2(doc, '5.3 Agent辅助写作流程')
    body(doc, 'Agent辅助写作是本系统的核心功能，采用Human-in-the-loop（人在回路）的协作模式。系统按照"文献解析→大纲生成→内容撰写→润色优化→引用检查→格式排版"的顺序依次调度6个Agent，每个Agent执行完成后等待用户审阅和确认。论文工作台页面如图5-4所示，展示了完整的6阶段流水线界面。')
    img(doc, '11_论文工作台_流水线.png', 12)
    cap(doc, '图5-4 论文工作台页面')

    body(doc, '论文工作台的主要操作流程如下：用户点击"开始写作"按钮启动全流程。系统首先调度文献解析Agent，该Agent读取用户关联的参考文献，利用大语言模型解析每篇文献的研究问题、方法和结论，输出结构化文献摘要。用户审阅摘要内容后选择确认或驳回。确认通过后，大纲生成Agent根据文献摘要和用户选择的论文模板生成结构化JSON大纲。用户确认大纲后，内容撰写Agent按大纲逐章撰写完整的论文正文。撰写完成后，润色Agent对论文进行语言流畅度提升和术语统一。随后引用检查Agent校验引用的准确性和格式规范。最后，格式排版Agent根据用户选择的格式模板（支持上传Word模板自动解析格式规则）对论文进行版面编排，输出带有排版标记的Markdown文本。')

    body(doc, '图5-5展示了论文工作台下半部分的流水线状态和内容预览区域。每个Agent的执行状态以不同颜色标签标识，已完成阶段为绿色，当前执行阶段为蓝色，待执行阶段为灰色。右侧面板实时显示Agent的流式输出内容。')
    img(doc, '12_论文工作台_流水线下半.png', 12)
    cap(doc, '图5-5 流水线状态与内容预览')

    body(doc, '论文内容预览与编辑区域如图5-6所示。用户可以在该区域查看已完成的论文全文，支持直接编辑修改。论文内容以Markdown格式呈现，标题层级清晰，引用标记以[N]格式标注。用户可通过"导出Word"按钮将论文导出为规范格式的Word文档。')
    img(doc, '13_论文工作台_内容区域.png', 12)
    cap(doc, '图5-6 论文内容预览与编辑')

    # 5.4
    h2(doc, '5.4 论文管理功能')
    body(doc, '论文列表页面展示了用户创建的所有论文，每篇论文以卡片形式展示论文标题、当前状态和创建时间。不同状态使用不同颜色标签区分：草稿为灰色、解析中为蓝色、大纲生成中为紫色、撰写中为黄色、润色中为粉色、检查中为橙色、排版中为青色、已完成为绿色。用户可通过"新建论文"按钮创建论文，创建时需输入主题、选择模板类型，并可勾选参考文献。论文列表页面如图5-7所示。')
    img(doc, '02_论文列表.png', 12)
    cap(doc, '图5-7 论文列表页面')

    # 5.5
    h2(doc, '5.5 文献库管理功能')
    body(doc, '文献库页面为用户提供个人学术文献的集中管理功能，如图5-8所示。用户可通过添加文献表单手动输入文献信息（标题、作者、出处、摘要、关键词等），也可上传PDF、Word或TXT格式文件自动解析文献全文内容。文献库以表格形式展示，每条文献包含标题、作者、出处、摘要等信息，支持查看详情和删除操作。')
    img(doc, '03_文献库.png', 12)
    cap(doc, '图5-8 文献库页面')

    # 5.6
    h2(doc, '5.6 个人中心')
    body(doc, '个人中心页面为用户提供个人信息管理和系统配置功能，如图5-9所示。页面包含四个功能区域：基本信息展示（用户名、邮箱、角色、注册时间）、修改资料（更新邮箱）、修改密码（需验证原密码后设置新密码）和DeepSeek API Key配置（用户可设置自己的API Key，不设置则使用系统默认Key）。页面底部还展示了用户最近创建的论文列表。')
    img(doc, '04_个人中心.png', 12)
    cap(doc, '图5-9 个人中心页面')

    # 5.7
    h2(doc, '5.7 管理后台')
    body(doc, '管理后台为管理员角色提供系统管理功能，包含四个标签页：概览、用户管理、论文管理和系统日志。概览页展示用户总数和论文总数的统计卡片，如图5-10所示。')
    img(doc, '05_管理后台_概览.png', 12)
    cap(doc, '图5-10 管理后台概览页')
    body(doc, '用户管理页支持查看所有用户列表、按用户名或邮箱搜索用户、提升/降级用户角色（普通用户↔管理员）和删除用户，如图5-11所示。')
    img(doc, '06_管理后台_用户管理.png', 12)
    cap(doc, '图5-11 管理后台用户管理页')
    body(doc, '论文管理页支持查看所有用户创建的论文并按状态筛选、删除异常论文，如图5-12所示。')
    img(doc, '07_管理后台_论文管理.png', 12)
    cap(doc, '图5-12 管理后台论文管理页')
    body(doc, '系统日志页展示Agent执行的详细日志记录，支持按级别筛选，便于监控系统运行状况和排查问题，如图5-13所示。')
    img(doc, '08_管理后台_日志.png', 12)
    cap(doc, '图5-13 管理后台系统日志页')

    out = os.path.join(BASE, 'PaperCraft_第5章_系统实现.docx')
    doc.save(out)
    print(f'第5章已保存: {out}')


# ==================== 第6章 ====================
def create_ch6():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')

    h1(doc, '第6章 系统测试')
    body(doc, '系统测试是软件开发过程中不可或缺的一个重要环节。它涉及到对完全集成的软件系统进行验证，以确认系统是否满足需求规格说明中定义的功能和性能要求。本章将对PaperCraft多智能体论文写作助手进行系统测试，包括测试环境说明、功能测试用例设计以及测试结果分析。')

    h2(doc, '6.1 测试意义')
    body(doc, '系统测试的意义在于通过系统化的测试流程，尽可能多地发现系统中存在的缺陷和问题，确保系统在上线前达到预期的质量和可靠性标准。对于本系统而言，测试的核心目标是验证各Agent模块的正确协作、用户认证的可靠性、数据操作的完整性以及前后端交互的流畅性。')

    h2(doc, '6.2 测试方法')
    body(doc, '本系统的测试采用黑盒测试与单元测试相结合的方式。黑盒测试不关注程序内部的结构和实现，只检验程序是否能够按照需求规格说明正常接收输入数据并产生正确的输出信息，是从用户观点出发的测试方法。单元测试则使用pytest框架对后端的核心模块进行独立测试，确保各模块的正确性和稳定性。')

    h2(doc, '6.3 测试环境')
    body(doc, '系统测试环境配置如下：')
    ttitle(doc, '表6-1 测试环境配置表')
    tbl(doc, ['项目', '配置'],
        [('操作系统', 'Windows 11 专业版'),
         ('CPU', 'Intel Core i7'),
         ('内存', '16GB'),
         ('后端框架', 'FastAPI 0.115.0'),
         ('前端框架', 'React 19 + TypeScript'),
         ('数据库', 'MySQL 8.0'),
         ('AI模型', 'DeepSeek Chat (via LiteLLM)'),
         ('自动化测试', 'Playwright (截图)'),
         ('单元测试', 'pytest 9.1.1 (后端) / Vitest (前端)')])

    h2(doc, '6.4 功能测试')
    body(doc, '功能测试主要针对系统的核心功能模块设计测试用例，包括用户注册与登录模块、论文管理模块、Agent写作模块、文献库管理模块和管理后台模块。')

    h3(doc, '6.4.1 注册与登录模块测试')
    body(doc, '注册与登录模块测试覆盖了用户注册和登录的各类场景，包括正常操作和异常操作。测试用例如表6-2所示。')
    ttitle(doc, '表6-2 注册与登录模块测试用例')
    tbl(doc, ['编号', '测试功能', '执行步骤', '期望结果', '测试结果'],
        [('TC01', '正常注册', '1.进入注册页 2.输入合法信息 3.点击注册', '注册成功，跳转首页', '通过'),
         ('TC02', '用户名过短', '1.输入1个字符的用户名 2.点击注册', '提示"用户名至少2个字符"', '通过'),
         ('TC03', '密码过短', '1.输入少于6位的密码 2.点击注册', '提示"密码至少6个字符"', '通过'),
         ('TC04', '用户名已存在', '1.使用已注册的用户名注册 2.点击注册', '提示"用户名已存在"', '通过'),
         ('TC05', '正常登录', '1.输入正确的账号密码 2.点击登录', '登录成功跳转首页', '通过'),
         ('TC06', '密码错误', '1.输入错误的密码 2.点击登录', '提示"用户名或密码错误"', '通过')])

    h3(doc, '6.4.2 论文管理模块测试')
    body(doc, '论文管理模块测试验证论文的创建、查看和删除功能是否正常。测试用例如表6-3所示。')
    ttitle(doc, '表6-3 论文管理模块测试用例')
    tbl(doc, ['编号', '测试功能', '执行步骤', '期望结果', '测试结果'],
        [('TC07', '创建论文', '1.点击新建论文 2.输入主题选择模板 3.创建', '论文出现在列表中', '通过'),
         ('TC08', '查看论文工作台', '1.点击论文卡片 2.进入工作台', '显示流水线和内容区', '通过'),
         ('TC09', '删除论文', '1.点击删除按钮 2.确认删除', '论文从列表移除', '通过')])

    h3(doc, '6.4.3 Agent协作流程测试')
    body(doc, 'Agent协作流程是本系统的核心功能，测试验证6个Agent的顺序执行和流式输出是否正常。测试用例如表6-4所示。')
    ttitle(doc, '表6-4 Agent协作流程测试用例')
    tbl(doc, ['编号', '测试功能', '执行步骤', '期望结果', '测试结果'],
        [('TC10', '一键全流程', '1.创建论文并关联文献 2.点击开始写作', '6个Agent依次执行', '通过'),
         ('TC11', 'SSE流式输出', '1.执行任意Agent 2.观察输出面板', '面板实时显示token输出', '通过'),
         ('TC12', '用户审阅反馈', '1.Agent执行完成 2.批准/驳回/编辑', '反馈处理后进下阶段', '通过'),
         ('TC13', '导出Word', '1.论文完成后点击导出 2.下载docx', '生成规范Word文档', '通过'),
         ('TC14', '格式模板上传', '1.上传docx模板 2.执行排版Agent', '按模板规则编排论文', '通过')])

    h3(doc, '6.4.4 管理后台模块测试')
    body(doc, '管理后台模块测试验证管理员特有的用户管理、论文管理和系统日志功能。测试用例如表6-5所示。')
    ttitle(doc, '表6-5 管理后台模块测试用例')
    tbl(doc, ['编号', '测试功能', '执行步骤', '期望结果', '测试结果'],
        [('TC15', '查看用户列表', '1.管理员登录 2.用户管理标签', '显示所有用户列表', '通过'),
         ('TC16', '搜索用户', '1.在搜索框输入关键字', '显示匹配的用户', '通过'),
         ('TC17', '提升为管理员', '1.点击"提升为管理员"', '用户角色变更为管理员', '通过'),
         ('TC18', '降级为用户', '1.点击"降级为用户"', '角色变更为普通用户', '通过'),
         ('TC19', '删除用户', '1.点击删除 2.确认', '用户及论文被删除', '通过'),
         ('TC20', '按状态筛选论文', '1.选择论文状态 2.查看结果', '显示符合条件的论文', '通过'),
         ('TC21', '查看系统日志', '1.进入系统日志标签', '显示Agent执行日志', '通过')])

    h2(doc, '6.5 单元测试')
    body(doc, '单元测试是对系统最小可测试单元（函数、方法、类）的验证。本系统后端使用pytest框架编写了28个单元测试用例，覆盖FormatAgent、FormatTemplate模型和格式模板API三个模块。测试结果如表6-6所示。')
    ttitle(doc, '表6-6 单元测试结果')
    tbl(doc, ['测试模块', '测试文件', '用例数', '通过数', '失败数'],
        [('FormatAgent', 'test_format_agent.py', '8', '8', '0'),
         ('FormatTemplate模型', 'test_format_template.py', '4', '4', '0'),
         ('格式模板API', 'test_format_template_api.py', '16', '16', '0'),
         ('合计', '-', '28', '28', '0')])
    body(doc, '所有28个单元测试全部通过，测试覆盖了Agent实例化、prompt格式化、模型字段定义、CRUD操作、文件上传验证、权限边界检查等多个维度，确保了系统核心功能的正确性和稳定性。')

    h2(doc, '6.6 测试结论')
    body(doc, '通过功能测试和单元测试的综合验证，PaperCraft多智能体论文写作助手的各功能模块均能正常运行。用户注册登录流程顺畅，论文管理和Agent写作功能完备，管理后台操作正常。28个后端单元测试全部通过，系统运行稳定，功能满足课程设计需求。')

    out = os.path.join(BASE, 'PaperCraft_第6章_系统测试.docx')
    doc.save(out)
    print(f'第6章已保存: {out}')


if __name__ == '__main__':
    create_ch5()
    create_ch6()
