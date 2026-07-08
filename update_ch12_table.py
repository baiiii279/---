#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""将第1-2章文档中Agent职责表的纯文本替换为规范Word表格"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(BASE, 'PaperCraft_第1-2章.docx')
OUT = os.path.join(BASE, 'PaperCraft_第1-2章_更新.docx')

def set_font(run, cn='宋体', en='Times New Roman', sz=Pt(10.5), bold=False):
    run.font.size = sz
    run.font.bold = bold
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn('w:eastAsia'), cn)

def cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_font(run, bold=bold)
    p.alignment = align

doc = Document(SRC)

# 找到"表2-2"段落位置
table_title_idx = -1
for i, para in enumerate(doc.paragraphs):
    if '表2-2' in para.text and 'Agent职责表' in para.text:
        table_title_idx = i
        break

if table_title_idx < 0:
    print('未找到 "表2-2 Agent职责表"')
    exit(1)

print(f'找到 "表2-2" 在段落[{table_title_idx}]')

# "表2-2" 后面应该紧跟着18个段落（表头3个 + 5行×3个 = 18个单元格）
# 但实际上文本里每个单元格占一个独立段落
# 解析这些单元格值
cells = []
for j in range(table_title_idx + 1, len(doc.paragraphs)):
    t = doc.paragraphs[j].text.strip()
    if not t:
        continue
    # 如果遇到"Agent协作流程"或"图2-2"或"2.2"等后续内容，停止
    if t.startswith('Agent协作流程') or t.startswith('图2-2') or t.startswith('2.2'):
        break
    cells.append(t)

print(f'找到 {len(cells)} 个单元格')

if len(cells) < 18:
    print(f'警告: 只找到 {len(cells)} 个单元格，需要18个')
    # 仍然尝试处理已有的

# 按3个一组分出行
rows_data = [cells[i:i+3] for i in range(0, len(cells), 3)]
print(f'分为 {len(rows_data)} 行')
for i, row in enumerate(rows_data):
    print(f'  行{i}: {row}')

# 先保存"表2-2"段落的引用
title_para = doc.paragraphs[table_title_idx]

# 删除原来的纯文本段落
first_data_idx = table_title_idx + 1

# 计算要删多少段：直到遇到后续内容
last_data_idx = first_data_idx
for j in range(first_data_idx, len(doc.paragraphs)):
    t = doc.paragraphs[j].text.strip()
    if t.startswith('Agent协作流程') or t.startswith('图2-2') or t.startswith('2.2'):
        break
    last_data_idx = j

# 从后往前删
for j in range(last_data_idx, first_data_idx - 1, -1):
    if j < len(doc.paragraphs):
        p = doc.paragraphs[j]
        p._element.getparent().remove(p._element)

# 创建表格
table = doc.add_table(rows=len(rows_data), cols=3, style='Table Grid')
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 填充数据
for r_idx, row in enumerate(rows_data):
    for c_idx, val in enumerate(row[:3]):  # 最多3列
        is_header = (r_idx == 0)
        cell_text(table.rows[r_idx].cells[c_idx], val, bold=is_header,
                  align=WD_ALIGN_PARAGRAPH.CENTER if c_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT)

# 列宽
for row in table.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(3.0)
    row.cells[2].width = Cm(9.5)

# 将表格插入到"表2-2"后面
title_para._element.addnext(table._element)
print('表格已插入')

doc.save(OUT)
print(f'文档已保存: {OUT}')
